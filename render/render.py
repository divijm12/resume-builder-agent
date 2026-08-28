#!/usr/bin/env python3
"""Stage 3 -- Render.

tailored_resume dict (the output of tailor.py, or its "tailored_resume" key
alone) in, resume + resume out under outputs/<company>_<role>_<date>/, named
<Firstname>_<Lastname>_<Company>_<Role>.docx/.pdf. One consistent template
(this file), not one-off formatting per application, per ARCHITECTURE.md.

Hard one-page rule: every bullet/skill/section the tailoring stage selected
must appear -- content is never dropped to force a one-page fit. Instead
margins are tightened first (free -- doesn't affect readability), then font
size/leading/spacing are scaled down, searching from most spacious to most
compact until the render is exactly one page. If even maximum compaction
still overflows, the most compact layout is used anyway and a warning is
printed -- overflowing by a fraction of a page beats losing content.

docx (python-docx) and pdf (reportlab) are two independent renders of the
same content -- not a docx->pdf conversion, since Word's AppleScript
automation for that is broken on this machine (Word rejects "save as" at
runtime with -1708 despite the command being in its own .sdef). Page-fit is
searched for using the PDF render as ground truth (reportlab reports exact
page count; docx has no accessible pagination info without an actual
rendering engine) and the winning (margin, scale) is applied to both --
so the docx fit is a close estimate calibrated off the PDF, not
independently verified. Both use Calibri -- docx by font-name reference
(relies on the reader having it installed); pdf by embedding the actual TTF
files bundled with this machine's Microsoft Word install, since a PDF needs
real font data, not just a name (falls back to Helvetica with a printed
warning if no Calibri file is found).

Section ordering (experience, projects, education, certifications) is
strictly reverse-chronological -- enforced in tailor.py by parsing each
entry's date field(s), not left to the model's own ordering choice.

Unlike agents/*.py, this stage does write files -- that's its job. Agents
stay pure (JSON in, JSON out); render/ and review/ are the persistence layer
one level up.
"""

import argparse
import json
import re
import sys
import tempfile
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT_NAME = "Calibri"
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)

# Most spacious first. Margin is tightened to its floor (0.4in) before font
# scale is touched at all, since smaller margins don't hurt readability the
# way smaller text does.
LAYOUT_CANDIDATES = [
    (0.6, 1.00), (0.5, 1.00), (0.4, 1.00),
    (0.4, 0.95), (0.4, 0.90), (0.4, 0.85), (0.4, 0.80),
]


MONTH_ABBR = {
    1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr", 5: "May", 6: "Jun",
    7: "Jul", 8: "Aug", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec",
}


def _format_month_year(date_str):
    """'2024-05' -> 'May 2024'; '2025' stays '2025' (no month to abbreviate);
    'present' -> 'Present'; anything else passes through unchanged."""
    if not date_str:
        return ""
    s = str(date_str).strip()
    if s.lower() == "present":
        return "Present"
    m = re.match(r"^(\d{4})-(\d{1,2})$", s)
    if m:
        year, month = int(m.group(1)), int(m.group(2))
        abbr = MONTH_ABBR.get(month)
        return f"{abbr} {year}" if abbr else s
    return s


def _format_dates(start, end):
    if not start:
        return ""
    start_label = _format_month_year(start)
    end_label = _format_month_year(end)
    return f"{start_label} -- {end_label}" if end_label else start_label


def _slug(s: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", s).strip("_")


# (regular, bold) TTF paths to try, in order -- bundled with Microsoft Word on
# this machine. reportlab needs an actual embedded font file for the PDF;
# unlike docx, a font *name* alone isn't enough since PDF viewers don't
# substitute installed system fonts the way Word does.
CALIBRI_CANDIDATES = [
    (
        "/Applications/Microsoft Word.app/Contents/Resources/DFonts/Calibri.ttf",
        "/Applications/Microsoft Word.app/Contents/Resources/DFonts/Calibrib.ttf",
    ),
    ("/Library/Fonts/Calibri.ttf", "/Library/Fonts/Calibrib.ttf"),
    (str(Path.home() / "Library/Fonts/Calibri.ttf"), str(Path.home() / "Library/Fonts/Calibrib.ttf")),
]

_calibri_registered = None  # cached (regular_font_name, bold_font_name) after first lookup


def _register_calibri_fonts():
    """Register real Calibri TTFs with reportlab if found on this machine, so
    the PDF uses actual Calibri glyphs. Falls back to Helvetica with a printed
    warning if no Calibri file is found -- never silently swaps fonts without
    saying so. docx is unaffected either way (font-by-name, not embedded)."""
    global _calibri_registered
    if _calibri_registered is not None:
        return _calibri_registered

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    for regular_path, bold_path in CALIBRI_CANDIDATES:
        if Path(regular_path).exists() and Path(bold_path).exists():
            pdfmetrics.registerFont(TTFont("Calibri", regular_path))
            pdfmetrics.registerFont(TTFont("Calibri-Bold", bold_path))
            _calibri_registered = ("Calibri", "Calibri-Bold")
            return _calibri_registered

    print(
        "Warning: no Calibri font file found on this machine -- PDF falling back to "
        "Helvetica. docx is unaffected (uses Calibri by name).",
        file=sys.stderr,
    )
    _calibri_registered = ("Helvetica", "Helvetica-Bold")
    return _calibri_registered


def _split_name(full_name: str):
    parts = full_name.strip().split()
    if not parts:
        return "Resume", ""
    return parts[0], "_".join(parts[1:])


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _set_font(run, size, bold=False, name=FONT_NAME):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold


def _add_section_heading(doc, text, sizes):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sizes["section_space_before"])
    p.paragraph_format.space_after = Pt(sizes["section_space_after"])
    run = p.add_run(text.upper())
    _set_font(run, size=Pt(sizes["section"]), bold=True)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def _add_bullet(doc, text, sizes):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_font(run, size=Pt(sizes["body"]))
    p.paragraph_format.space_after = Pt(sizes["bullet_space_after"])


def _add_heading_row(doc, left_text, right_text, sizes, right_tab):
    """Bold left-aligned text with a right-aligned date/label via a tab stop."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sizes["row_space_before"])
    p.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(left_text)
    _set_font(run, size=Pt(sizes["body"]), bold=True)
    if right_text:
        date_run = p.add_run(f"\t{right_text}")
        _set_font(date_run, size=Pt(sizes["small"]))


def _docx_sizes(scale: float) -> dict:
    return {
        "name": 20 * scale,
        "contact": 9.5 * scale,
        "section": 12 * scale,
        "body": 10.5 * scale,
        "small": 9.5 * scale,
        "section_space_before": 10 * scale,
        "section_space_after": 2 * scale,
        "row_space_before": 6 * scale,
        "bullet_space_after": 4 * scale,
        "contact_space_after": 6 * scale,
    }


def render_docx(tailored_resume: dict, output_path: Path, margin_in: float = 0.4, scale: float = 1.0) -> None:
    doc = Document()
    sizes = _docx_sizes(scale)
    margin = Inches(margin_in)
    right_tab = PAGE_WIDTH - 2 * margin

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(sizes["body"])
    for section in doc.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = Inches(max(0.3, margin_in - 0.1))
        section.bottom_margin = Inches(max(0.3, margin_in - 0.1))

    basics = tailored_resume.get("basics", {})

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(basics.get("name", ""))
    _set_font(run, size=Pt(sizes["name"]), bold=True)
    p.paragraph_format.space_after = Pt(2 * scale)

    contact_parts = [x for x in [basics.get("phone"), basics.get("email")] if x]
    links = basics.get("links", {}) or {}
    for key in ("linkedin", "github", "portfolio"):
        if links.get(key):
            contact_parts.append(links[key])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(" | ".join(contact_parts))
    _set_font(run, size=Pt(sizes["contact"]))
    p.paragraph_format.space_after = Pt(sizes["contact_space_after"])

    if tailored_resume.get("education"):
        _add_section_heading(doc, "Education", sizes)
        for edu in tailored_resume["education"]:
            _add_heading_row(doc, f"{edu.get('degree', '')} -- {edu.get('institution', '')}", "", sizes, right_tab)
            if edu.get("honors"):
                p = doc.add_paragraph()
                run = p.add_run(edu["honors"])
                _set_font(run, size=Pt(sizes["small"]))
                p.paragraph_format.space_after = Pt(sizes["bullet_space_after"])

    if tailored_resume.get("certifications"):
        _add_section_heading(doc, "Certifications", sizes)
        for cert in tailored_resume["certifications"]:
            year = f" ({cert['year']})" if cert.get("year") else ""
            _add_bullet(doc, f"{cert.get('name', '')}{year}", sizes)

    if tailored_resume.get("skills"):
        _add_section_heading(doc, "Skills", sizes)
        p = doc.add_paragraph()
        run = p.add_run(", ".join(tailored_resume["skills"]))
        _set_font(run, size=Pt(sizes["body"]))

    if tailored_resume.get("experience"):
        _add_section_heading(doc, "Experience", sizes)
        for exp in tailored_resume["experience"]:
            left = f"{exp.get('title', '')} -- {exp.get('company', '')}"
            _add_heading_row(doc, left, _format_dates(exp.get("start"), exp.get("end")), sizes, right_tab)
            for b in exp.get("bullets", []):
                _add_bullet(doc, b["text"], sizes)

    if tailored_resume.get("projects"):
        _add_section_heading(doc, "Projects", sizes)
        for proj in tailored_resume["projects"]:
            tech = " -- " + " | ".join(proj["tech"]) if proj.get("tech") else ""
            left = f"{proj.get('name', '')}{tech}"
            _add_heading_row(doc, left, _format_month_year(proj.get("date", "")), sizes, right_tab)
            for b in proj.get("bullets", []):
                _add_bullet(doc, b["text"], sizes)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def render_pdf(tailored_resume: dict, output_path: Path, margin_in: float = 0.4, scale: float = 1.0) -> None:
    from xml.sax.saxutils import escape

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Table,
        TableStyle,
    )

    def sz(base):
        return base * scale

    regular_font, bold_font = _register_calibri_fonts()

    styles = {
        "name": ParagraphStyle("name", fontName=bold_font, fontSize=sz(20), leading=sz(24), alignment=TA_CENTER, spaceAfter=2 * scale),
        "contact": ParagraphStyle("contact", fontName=regular_font, fontSize=sz(9.5), leading=sz(12), alignment=TA_CENTER, spaceAfter=6 * scale),
        "section": ParagraphStyle("section", fontName=bold_font, fontSize=sz(12), leading=sz(14), spaceBefore=10 * scale, spaceAfter=2 * scale),
        "body": ParagraphStyle("body", fontName=regular_font, fontSize=sz(10.5), leading=sz(13)),
        "bold_row": ParagraphStyle("bold_row", fontName=bold_font, fontSize=sz(10.5), leading=sz(13)),
        "date_row": ParagraphStyle("date_row", fontName=regular_font, fontSize=sz(9.5), leading=sz(13), alignment=TA_RIGHT),
        "bullet": ParagraphStyle("bullet", fontName=regular_font, fontSize=sz(10.5), leading=sz(13)),
        "small": ParagraphStyle("small", fontName=regular_font, fontSize=sz(9.5), leading=sz(12), spaceAfter=3 * scale),
    }

    story = []
    basics = tailored_resume.get("basics", {})
    story.append(Paragraph(escape(basics.get("name", "")), styles["name"]))
    contact_parts = [x for x in [basics.get("phone"), basics.get("email")] if x]
    links = basics.get("links", {}) or {}
    for key in ("linkedin", "github", "portfolio"):
        if links.get(key):
            contact_parts.append(links[key])
    story.append(Paragraph(escape(" | ".join(contact_parts)), styles["contact"]))

    def section_heading(text):
        story.append(Paragraph(text.upper(), styles["section"]))
        story.append(HRFlowable(width="100%", thickness=0.75, color=colors.black, spaceAfter=4 * scale))

    def heading_row(left, right):
        usable = PAGE_WIDTH / 914400 - 2 * margin_in  # EMU->in, minus margins
        t = Table(
            [[Paragraph(escape(left), styles["bold_row"]), Paragraph(escape(right), styles["date_row"])]],
            colWidths=[(usable - 1.7) * inch, 1.7 * inch],
        )
        # reportlab Tables default to hAlign='CENTER' -- any tiny mismatch between
        # `usable` above and the frame's actual computed width then shifts the
        # whole table left/right relative to plain (non-Table) paragraphs like
        # section headings and bullets, which always start flush at the frame
        # edge. Pin it explicitly so every heading_row lines up with everything
        # else regardless of that rounding, for any resume's content.
        t.hAlign = "LEFT"
        t.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 4 * scale),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        story.append(t)

    def bullet_list(items):
        # bulletFontSize must match the paragraph's own font size -- reportlab
        # does not inherit it from the Paragraph style, and left at its
        # unrelated default it draws the dot using the wrong font metrics,
        # which is barely visible on a one-line item but drifts further from
        # the first line's baseline the more lines an item wraps to. This is
        # a general correctness fix (derived from the same dynamic `scale`
        # used everywhere else), not tuned to any one resume's content.
        story.append(
            ListFlowable(
                [
                    ListItem(
                        Paragraph(escape(b), styles["bullet"]),
                        spaceAfter=3 * scale,
                        bulletFontName=regular_font,
                        bulletFontSize=styles["bullet"].fontSize,
                    )
                    for b in items
                ],
                bulletType="bullet",
                leftIndent=14,
            )
        )

    if tailored_resume.get("education"):
        section_heading("Education")
        for edu in tailored_resume["education"]:
            heading_row(f"{edu.get('degree', '')} -- {edu.get('institution', '')}", "")
            if edu.get("honors"):
                story.append(Paragraph(escape(edu["honors"]), styles["small"]))

    if tailored_resume.get("certifications"):
        section_heading("Certifications")
        cert_lines = [
            f"{c.get('name', '')}" + (f" ({c['year']})" if c.get("year") else "")
            for c in tailored_resume["certifications"]
        ]
        bullet_list(cert_lines)

    if tailored_resume.get("skills"):
        section_heading("Skills")
        story.append(Paragraph(escape(", ".join(tailored_resume["skills"])), styles["body"]))

    if tailored_resume.get("experience"):
        section_heading("Experience")
        for exp in tailored_resume["experience"]:
            heading_row(
                f"{exp.get('title', '')} -- {exp.get('company', '')}",
                _format_dates(exp.get("start"), exp.get("end")),
            )
            bullet_list([b["text"] for b in exp.get("bullets", [])])

    if tailored_resume.get("projects"):
        section_heading("Projects")
        for proj in tailored_resume["projects"]:
            tech = " -- " + " | ".join(proj["tech"]) if proj.get("tech") else ""
            heading_row(f"{proj.get('name', '')}{tech}", _format_month_year(proj.get("date", "")))
            bullet_list([b["text"] for b in proj.get("bullets", [])])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=margin_in * inch,
        rightMargin=margin_in * inch,
        topMargin=max(0.3, margin_in - 0.1) * inch,
        bottomMargin=max(0.3, margin_in - 0.1) * inch,
    )
    doc.build(story)


# ---------------------------------------------------------------------------
# One-page layout search
# ---------------------------------------------------------------------------

def find_one_page_layout(tailored_resume: dict) -> dict:
    """Find the most spacious (margin, scale) that fits the PDF render onto
    exactly one page, without dropping any content. Tightens margins to
    their floor before touching font scale at all. If even the most compact
    candidate still overflows, returns it anyway with fits=False -- the
    caller must warn, not truncate content to force a fit."""
    from pypdf import PdfReader

    with tempfile.TemporaryDirectory() as tmp:
        probe_path = Path(tmp) / "probe.pdf"
        for margin_in, scale in LAYOUT_CANDIDATES:
            render_pdf(tailored_resume, probe_path, margin_in=margin_in, scale=scale)
            if len(PdfReader(str(probe_path)).pages) == 1:
                return {"margin_in": margin_in, "scale": scale, "fits": True}
        margin_in, scale = LAYOUT_CANDIDATES[-1]
        return {"margin_in": margin_in, "scale": scale, "fits": False}


def main():
    parser = argparse.ArgumentParser(description="Render a tailored resume to .docx and .pdf.")
    parser.add_argument("--tailored-json", required=True, help="Path to tailor.py's output JSON")
    parser.add_argument("--company", required=True, help="Company name, for the output folder/filename")
    parser.add_argument("--role", required=True, help="Role title, for the output folder/filename")
    parser.add_argument("--outputs-dir", type=Path, default=Path("outputs"), help="Base outputs directory")
    args = parser.parse_args()

    data = json.loads(Path(args.tailored_json).read_text())
    tailored_resume = data.get("tailored_resume", data)

    folder_name = f"{_slug(args.company)}_{_slug(args.role)}_{date.today().isoformat()}"
    output_dir = args.outputs_dir / folder_name

    if output_dir.exists():
        print(f"Error: {output_dir} already exists -- refusing to overwrite a previous version.", file=sys.stderr)
        sys.exit(1)

    first, last = _split_name(tailored_resume.get("basics", {}).get("name", ""))
    name_part = f"{_slug(first)}_{_slug(last)}" if last else _slug(first)
    file_base = f"{name_part}_{_slug(args.company)}_{_slug(args.role)}"

    layout = find_one_page_layout(tailored_resume)
    if not layout["fits"]:
        print(
            "Warning: content still exceeds one page even at maximum compaction "
            f"(margin={layout['margin_in']}in, font scale={layout['scale']}). "
            "Using the most compact layout anyway -- no content was dropped to force a fit.",
            file=sys.stderr,
        )

    docx_path = output_dir / f"{file_base}.docx"
    pdf_path = output_dir / f"{file_base}.pdf"
    render_docx(tailored_resume, docx_path, margin_in=layout["margin_in"], scale=layout["scale"])
    render_pdf(tailored_resume, pdf_path, margin_in=layout["margin_in"], scale=layout["scale"])

    print(json.dumps({"docx_path": str(docx_path), "pdf_path": str(pdf_path), "layout": layout}, indent=2))


if __name__ == "__main__":
    main()
